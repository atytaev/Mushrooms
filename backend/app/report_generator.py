import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

logger = logging.getLogger(__name__)
from .models import (
    Inspection,
    MushroomStorage,
    MushroomPhoto,
    ProductMarkingPhoto,
    QuantityInspection,
    QuantityInspectionPhoto,
    QualityInspection,
    QualityInspectionPhoto,
    DiameterMeasurement,
    DiameterMeasurementPhoto,
    Pallet,
    PalletPhoto,
    ProductLoading,
    ProductLoadingPhoto,
    Box,
    Thermometer,
)


def generate_inspection_report(inspection_id,
                               placement_ids=None,
                               marking_ids=None,
                               quantity_photo_ids=None,
                               quality_photo_ids=None,
                               pallet_photo_ids=None,
                               loading_ids=None):
    try:
        insp = Inspection.objects.get(id=inspection_id)
        qual = QualityInspection.objects.filter(inspection=insp).first()
        logger.debug(
            f"generate_inspection_report: placement_ids={placement_ids}, marking_ids={marking_ids}, loading_ids={loading_ids}")

        # 1. Размещение товара
        # Получаем единственный объект хранения (или None)
        storage_obj = MushroomStorage.objects.filter(inspection=insp).first()

        # Собираем фотографии только для него
        placement_photos = MushroomPhoto.objects.filter(
            id__in=placement_ids or [],
            storage__inspection_id=inspection_id
        )
        logger.debug(
            "Found %d placement_photos (IDs %r)",
            placement_photos.count(),
            list(placement_photos.values_list('id', flat=True))
        )

        # 2. Маркировка
        marking_photos = ProductMarkingPhoto.objects.filter(
            inspection=insp,
            id__in=marking_ids or []
        ) if marking_ids else []

        # 3. Инспекция количества товара
        qty = QuantityInspection.objects.filter(inspection=insp).first()
        inspection_rows = Box.objects.filter(quantity_inspection=qty)
        box = Box.objects.filter(quantity_inspection=qty).first()
        qty_photos = QuantityInspectionPhoto.objects.filter(
            quantity_inspection=qty
        )

        # 4. Инспекция качества товара
        qual = QualityInspection.objects.filter(inspection=insp).first()
        qual_photos = QualityInspectionPhoto.objects.filter(
            quality_inspection=qual
        )

        # 5. Замер диаметра грибов
        diam = DiameterMeasurement.objects.filter(inspection=insp).first()
        diam_photos = DiameterMeasurementPhoto.objects.filter(
            diameter_measurement=diam
        )

        # 6. Паллеты
        pallets = Pallet.objects.filter(inspection=insp)
        pallet_photos = []
        for p in pallets:
            pallet_photos += list(PalletPhoto.objects.filter(pallet=p))

        # 7. Погрузка
        loading = ProductLoading.objects.filter(inspection=insp).first()
        loading_photos = ProductLoadingPhoto.objects.filter(
            loading=loading
        )

        # --- формируем DOCX ---
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        section = doc.sections[0]
        section.top_margin = Inches(0.5)  # Отступ сверху (1 дюйм)
        section.bottom_margin = Inches(0.5)  # Отступ снизу (1 дюйм)
        section.left_margin = Inches(0.5)  # Отступ слева (1 дюйм)
        section.right_margin = Inches(0.5)  # Отступ справа (1 дюйм)

        # Страница 1: Заголовок и информация
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(100)
        run = paragraph.add_run(f"Номер работы: {insp.job_number}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Цвет шрифта

        paragraph = doc.add_paragraph()
        run = paragraph.add_run("В соответствии с заявкой, полученной от нашего Заказчика")
        run.font.size = Pt(12)

        # Вставка оставшегося текста (не изменяем стиль)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("ТОО «ЭКО ГОРОД ТРЭЙД», РЕСПУБЛИКА КАЗАХСТАН")
        run.bold = True  # Делаем текст жирным
        run.font.size = Pt(12)  # Уменьшение размера шрифта
        paragraph.alignment = 1  # Выравнивание по центру

        paragraph = doc.add_paragraph()
        run = paragraph.add_run("и согласно следующей инструкции:")
        run.font.size = Pt(12)

        doc.add_paragraph(
            "ИНСПЕКЦИЯ КОЛИЧЕСТВА И ВИЗУАЛЬНАЯ ИНСПЕКЦИЯ КАЧЕСТВА ТОВАРА\n"
            "СОГЛАСНО ИНСТРУКЦИЯМ КЛИЕНТА\n\n"
        ).alignment = 1  # по центру

        # Создаем таблицу для товара и количества
        table = doc.add_table(rows=1, cols=2)

        # Устанавливаем ширину столбцов
        table.columns[0].width = Pt(200)
        table.columns[1].width = Pt(500)

        # Первая ячейка
        cell1 = table.cell(0, 0)
        cell1.text = "ТОВАР И КОЛИЧЕСТВО ЗАЯВЛЕНО КАК:"
        cell1.paragraphs[0].runs[0].font.size = Pt(12)

        # Вторая ячейка
        quantity_of_boxes = storage_obj.quantity_of_boxes

        cell2 = table.cell(0, 1)
        cell2.text = f"ШАМПИНЬОНЫ СВЕЖИЕ КУЛЬТИВИРУЕМЫЕ\nВЫСШИЙ СОРТ, {quantity_of_boxes} ЯЩИКОВ\n"
        cell2.paragraphs[0].runs[0].font.size = Pt(12)

        cell1.paragraphs[0].alignment = 0
        cell2.paragraphs[0].alignment = 0

        # Добавляем новую строку в таблицу
        row = table.add_row()

        # Заполняем новую строку
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        cell1.text = "НОМЕР СЧЕТ-ФАКТУРЫ:"
        cell2.text = f"{storage_obj.invoice_number}"

        # Добавляем отступ между строками таблицы
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(24)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(24)


        doc.add_paragraph(
            f"\nМЫ ПРОВЕЛИ ИНСПЕКЦИЮ И НАСТОЯЩИМ СООБЩАЕМ СЛЕДУЮЩЕЕ:"
        ).alignment = 0  # по левому краю


        # Создаем таблицу для места и даты инспекции
        table = doc.add_table(rows=1, cols=2)

        # Устанавливаем ширину столбцов
        table.columns[0].width = Pt(200)
        table.columns[1].width = Pt(1000)

        # Первая ячейка
        cell1 = table.cell(0, 0)
        cell1.text = "МЕСТО ИНСПЕКЦИИ:"
        cell1.paragraphs[0].runs[0].font.size = Pt(12)

        # Вторая ячейка
        cell2 = table.cell(0, 1)
        cell2.text = f"КФХ «ГРИБНАЯ СТРАНА», РЕСПУБЛИКА БЕЛАРУСЬ, БРЕСТСКАЯ ОБЛ., БАРАНОВИЧСКИЙ Р-Н, МАЛАХОВЕЦКИЙ С/С, ЗДАНИЕ 21 "
        cell2.paragraphs[0].runs[0].font.size = Pt(12)

        cell1.paragraphs[0].alignment = 0
        cell2.paragraphs[0].alignment = 0

        # Добавляем новую строку для даты
        row = table.add_row()

        # Заполняем новую строку
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        cell1.text = "ДАТА ИНСПЕЦИИ:"
        # Получаем сегодняшнюю дату и форматируем, например, в виде DD.MM.YYYY
        today_str = date.today().strftime("%d.%m.%Y")
        cell2.text = today_str

        paragraph = doc.add_paragraph()
        run = paragraph.add_run("\n     1. РАЗМЕЩЕНИЕ ТОВАРА\n\n")
        run.bold = True

        storage_obj = MushroomStorage.objects.filter(inspection=insp).first()
        if storage_obj.thermometer_id:
            try:
                thermometer_obj = Thermometer.objects.get(id=storage_obj.thermometer_id)
            except Thermometer.DoesNotExist:
                thermometer_obj = None

        thermometer_info = thermometer_obj.info if thermometer_obj else "-"
        thermometer_serial = thermometer_obj.serial if thermometer_obj else "-"
        thermometer_calibration = thermometer_obj.calibration_date.strftime('%d.%m.%Y') if (

        thermometer_obj and thermometer_obj.calibration_date) else "-"
        p = doc.add_paragraph(
            f"На момент проведения инспекции шампиньоны свежие, культивируемые высший сорт "
            f"(далее — «грибы») находились на хранении в промышленном холодильнике по {quantity_of_boxes/storage_obj.quantity_of_pallets} ящиков. "
            f"{storage_obj.quantity_of_pallets} поддона шампиньоны свежие, культивируемые высший сорт с калибром 50-70 мм.\n"
            f"Температура в холодильнике +{storage_obj.temperature_in_fridge}°C по Цельсию. Температура гриба в переделах от "
            f"{storage_obj.mushroom_temperature_min}°C до +{storage_obj.mushroom_temperature_max}°C по Цельсию."
            f" Измерение грибов производилось термометром электронным {thermometer_info},"
            f" идентификационный {thermometer_serial} (дата проверки - {thermometer_calibration})\n\n"
            )
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if placement_photos.exists():
            # page break + заголовок
            doc.add_page_break()

            cols = 2
            rows = (placement_photos.count() + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            table.autofit = True

            for idx, photo in enumerate(placement_photos):
                r, c = divmod(idx, cols)
                cell = table.cell(r, c)
                run = cell.paragraphs[0].add_run()
                run.add_picture(photo.image.path, width=Inches(4))
        else:
            doc.add_page_break()
            doc.add_paragraph("Пользователь не выбрал ни одного фото для отчёта.")
        logger.debug(f"Found {placement_photos.count()} placement photos for storage id={storage_obj.id if storage_obj else None}")

        # ПЕРЕХОД К РАЗДЕЛУ 2
        doc.add_page_break()
        p2 = doc.add_paragraph("    2. МАРКИРОВКА")
        p2.runs[0].bold = True
        doc.add_paragraph()  # пустая строка
        doc.add_paragraph(
            "Грибы были упакованы в полипропиленовые ящики. "
            "На ящиках с грибами имелись бумажные этикетки со следующей маркировкой: "
            "на первой этикетке информация содержит информационный характер продукта "
            "(наименование товара, производитель, импортер, срок хранения дата упаковки и т. д.), "
            "на второй этикетке информация для внутреннего контроля и идентификации товара и персонала "
            "(упаковщика, контролера и т. п.)."
        )

        if marking_photos:
            photos = ProductMarkingPhoto.objects.filter(id__in=marking_photos)
            if photos.exists():
                # опять таблица, по 2 на строку (или 3, как вам нужно)
                cols = 2
                rows = (photos.count() + cols - 1) // cols
                table = doc.add_table(rows=rows, cols=cols)
                for idx, photo in enumerate(photos):
                    r, c = divmod(idx, cols)
                    cell = table.cell(r, c)
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(photo.image.path, width=Inches(4))

        # раздел 3
        doc.add_page_break()
        p3 = doc.add_paragraph("    3. ИНСПЕКЦИЯ КОЛИЧЕСТВА ТОВАРА")
        p3.runs[0].bold = True
        doc.add_paragraph()

        scale = box.quantity_inspection.scale
        if scale:
            scale_model = scale.model
            calibration_date = scale.calibration_date.strftime('%d.%m.%Y') if scale.calibration_date else "-"
            serial_number = scale.serial_number or "-"
        else:
            scale_model = "-"
            calibration_date = "-"
            serial_number = "-"

        doc.add_paragraph(
            f"К инспекции было предоставлено {storage_obj.quantity_of_pallets} деревянных поддона по {quantity_of_boxes/storage_obj.quantity_of_pallets} ящиков на каждом поддоне. "
            f"Итого {quantity_of_boxes} ящиков с грибами. Для проверки качества грибов на соответствие, заявленного "
            f"производителем была произведена случайным образом выборка (заранее согласованная с клиентом) "
            f"по одному случайно выбранному ящику с каждого поддона. Выборка составила {storage_obj.quantity_of_pallets} ящика. Вес каждого "
            f"ящика с грибами в среднем не менее 3 кг. Взвешивание проводилось на весах электронных марки "
            f"{scale_model} дата поверки {calibration_date},"
            f" заводской номер {serial_number}.\n\n"
        )

          # --- фотографии инспекции количества ---
        qty_photos = QuantityInspectionPhoto.objects.filter(
            quantity_inspection = qty,
            id__in = (quantity_photo_ids or [])
        )
        if qty_photos.exists():
            cols = 2
            rows = (qty_photos.count() + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            for idx, photo in enumerate(qty_photos):
                r, c = divmod(idx, cols)
                cell = table.cell(r, c)
                cell.paragraphs[0].add_run().add_picture(photo.image.path, width=Inches(4))
        else:
            doc.add_paragraph("Пользователь не выбрал фото для раздела 3.")

        # раздел 4
        conforms_kg = float(qual.conforms_to_declared_grade)
        sample_mass = float(qual.sample_mass_kg)
        conforms_pct = conforms_kg / sample_mass * 100
        con_pct =  (sample_mass-conforms_kg)/sample_mass * 100


        doc.add_page_break()
        p4 = doc.add_paragraph("    4. ИНСПЕКЦИЯ КАЧЕСТВА ТОВАРА ")
        p4.runs[0].bold = True
        doc.add_paragraph()
        doc.add_paragraph(
            f"Объём выборки для визуальной инспекции качества товара (высший сорт, калибр 50–70 мм) "
            f"был произведен согласно инструкциям клиента и составил {storage_obj.quantity_of_pallets} ящика.\n\n"
            f"Было проверено {sample_mass} (объединённая проба) грибов по различным параметрам:\n"
            f"- внешний вид (цвет, загрязненность, целостность и т. п.);\n"
            f"- запах;\n"
            f"- калибр\n\n"
            f"Детали выборки следующие\n"
        )

        # создаём таблицу 1×3
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # заголовки
        hdr_cells = table.rows[0].cells
        headers = [
            "Фактическое количество товара,\nпредставленного к инспекции",
            "Объём выборки\n(количество случайным\nобразом отобранных ящиков)",
            "Масса объединённой пробы, кг"
        ]
        for cell, text in zip(hdr_cells, headers):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            p.alignment = 1  # по центру

        # добавляем строку с данными
        row_cells = table.add_row().cells
        values = [f"{quantity_of_boxes}", f"{storage_obj.quantity_of_pallets}", f"{sample_mass}"]
        for cell, val in zip(row_cells, values):
            p = cell.paragraphs[0]
            p.add_run(val)
            p.alignment = 1  # по центру



        doc.add_paragraph(
            "\n\nОпределение соответствия грибов из выборки высшему сорту (внешний вид, запах):\n"
            f"Соответствие заявленному сорту – {conforms_kg:.3f} кг или {conforms_pct:.2f}%\n"
            f"Несоответствие заявленному сорту*- {(sample_mass-conforms_kg):.3f} кг или {con_pct:.2f}\n"
            f"* несоответствие грибов: механические повреждения в виде надломов шляпки гриба,"
            f"в виде трещин и пустот ножки, потемневшие ножки, раскрытие шляпки и т. п.\n"
        )

        for run in p.runs:
            text = run.text
            if text.startswith("* несоответствие"):
                run.font.size = Pt(8)

        if qual.off_grade_mass_kg_50:
            calib_pct = float(qual.off_grade_mass_kg_50) / float(qual.sample_mass_kg) * 100
            doc.add_paragraph(
                f"При проверке также было обнаружено не соответствующие по калибру грибы:\n"
                f"-{qual.off_grade_mass_kg_50:.3f} кг - калибром менее 50мм."
                f"В процентном соотношении от объединенной пробы  {calib_pct:.2f}%"
            )
        if qual.off_grade_mass_kg_70:
            calib_pct = float(qual.off_grade_mass_kg_70) / float(qual.sample_mass_kg) * 100
            doc.add_paragraph(
                f"-{qual.off_grade_mass_kg_70:.3f} кг - калибром более 70мм."
                f"В процентном соотношении от объединенной пробы  {calib_pct:.2f}%"
            )

        qual = QualityInspection.objects.filter(inspection=insp).first()

        # получаем только те фото, которые пришли в запросе
        selected_qual_photos = QualityInspectionPhoto.objects.filter(
            quality_inspection=qual,
            id__in=quality_photo_ids or []
        )

        if selected_qual_photos.exists():
            cols = 2
            rows = (selected_qual_photos.count() + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            for idx, photo in enumerate(selected_qual_photos):
                r, c = divmod(idx, cols)
                cell = table.cell(r, c)
                run = cell.paragraphs[0].add_run()
                run.add_picture(photo.image.path, width=Inches(4))
        else:
            doc.add_paragraph("Пользователь не выбрал ни одного фото для раздела 4.")

        doc.add_paragraph(
            f"Было взвешено {storage_obj.quantity_of_pallets} поддона с товаром. "
            "Взвешивание проводилось на стационарных весах склада «А» (дата поверки апрель 2024г.) КФХ «Грибная страна»"
            "Данные по взвешиванию поддонов с товаром приведены в таблице:"
        )

        # формируем строки из inspectionRows
        rows = []
        for i, row in enumerate(inspection_rows, start=1):
            gross = round(row.net_weight)
            pallet_w = round(row.defect_weight)
            box_cnt = int(quantity_of_boxes/storage_obj.quantity_of_pallets)  # количество ящиков
            net = f"{gross-pallet_w-36.6}"  # вес нетто
            grade = "Обычный белый 50-70мм"
            rows.append((i, gross, pallet_w, box_cnt, net, grade))

        # Создаём таблицу 1+len(rows) × 6
        table = doc.add_table(rows=1 + len(rows), cols=6, style='Table Grid')

        tbl = table._tbl
        first_tr = tbl.tr_lst[0]
        trPr = first_tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)

        # 2) запрет разрыва строк
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

        # теперь заполняем заголовки
        hdr = table.rows[0].cells
        titles = ["№ п/п", "Вес брутто с поддоном, кг", "Вес поддона, кг", "Количество ящиков, кг", "Вес нетто, кг", "Вид и калибр, кг"]
        for cell, title in zip(hdr, titles):
            p = cell.paragraphs[0]
            run = p.add_run(title)
            run.bold = True
            p.alignment = 1

        # Заполняем данные
        for row_idx, data in enumerate(rows, start=1):
            cells = table.rows[row_idx].cells
            for col_idx, val in enumerate(data):
                p = cells[col_idx].paragraphs[0]
                p.add_run(str(val))
                p.alignment = 1

        # --- Считаем итоги по каждому числовому столбцу ---
        sum_gross = sum(r[1] for r in rows)
        sum_pallet = sum(r[2] for r in rows)
        sum_boxes = sum(r[3] for r in rows)
        # если net у вас строка, приведите к float/int:
        sum_net = sum(float(r[4]) for r in rows)

        # Добавляем строку «Итого»
        total_cells = table.add_row().cells
        titles = ["Итого", sum_gross, sum_pallet, sum_boxes, sum_net, ""]  # последний столбец пустой
        for idx, val in enumerate(titles):
            p = total_cells[idx].paragraphs[0]
            run = p.add_run(str(val))
            run.bold = True
            p.alignment = 1

        doc.add_paragraph(
            f"\nВес одного пустого ящика в среднем 210г."
            f"Вес пустых ящиков на одном поддоне плюс упаковка (в виде картонных упаковок 3 кг) в среднем 36,60 кг."
            f"В результате взвешивания получили. Вес брутто – {sum_gross} кг. Вес нетто – {sum_net} кг."
        )

        selected_pallet_photos = PalletPhoto.objects.filter(
            pallet__inspection=insp,
            id__in=pallet_photo_ids or []
        )
        if selected_pallet_photos.exists():
            cols = 2
            rows = (selected_pallet_photos.count() + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            for idx, photo in enumerate(selected_pallet_photos):
                r, c = divmod(idx, cols)
                cell = table.cell(r, c)
                run.bold = True
                # само изображение
                cell.paragraphs[0].add_run().add_picture(photo.image.path, width=Inches(4))
        else:
            doc.add_paragraph("Пользователь не выбрал ни одной фотографии палет.")

        if loading and loading.thermometer:
            thermometer_model = loading.thermometer.info or "-"
            thermometer_serial = loading.thermometer.serial or "-"
            thermometer_calibration = loading.thermometer.calibration_date.strftime(
                '%d.%m.%Y') if loading.thermometer.calibration_date else "-"
        else:
            thermometer_model = "-"
            thermometer_serial = "-"
            thermometer_calibration = "-"

        doc.add_paragraph(
            f"Во время инспекции и после отгрузки случайным образом были выбраны паллеты с грибами и измерена температура гриба."
            f" Температура гриба при загрузке в авторефрижератор составила +2,7 градуса по Цельсию."
            f"Измерения грибов производились термометром электронным {thermometer_model}, "
            f"идентификационный номер №{thermometer_serial} (дата поверки – {thermometer_calibration})."
        )

        if loading and loading.car_number:
            seal_number = loading.seal_number or "-"
            refrigerator_number = loading.refrigerator_number or "-"
            transport_temperature = loading.transport_temperature

            text = (
                f"К инспекции была представлена машина (гос. номер {loading.car_number}) "
                f"с авторефрижератором (гос. номер {refrigerator_number}). "
                "На момент инспекции он находился в удовлетворительном состоянии, без видимых повреждений, "
                "сквозных отверстий, чистый, сухой, без посторонних запахов. После погрузки всех поддонов "
                "с грибами авторефрижератор был опломбирован инспектором СЖС пломбой SGS "
                f"{seal_number}."
            )

            if transport_temperature is not None:
                text += f" Водителем была выставлена температура для транспортировки +{transport_temperature:.1f} градуса по Цельсию."

            doc.add_paragraph(text)
        else:
            # Можно и не добавлять ничего или добавить предупреждение
            pass

        selected_loading_photos = ProductLoadingPhoto.objects.filter(
            loading=loading,
            id__in=loading_ids or []
        )
        if selected_loading_photos.exists():
            cols = 2
            rows = (selected_loading_photos.count() + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            table.autofit = True
            for idx, photo in enumerate(selected_loading_photos):
                r, c = divmod(idx, cols)
                cell = table.cell(r, c)
                # само изображение
                cell.paragraphs[0].add_run().add_picture(photo.image.path, width=Inches(4))
        else:
            doc.add_paragraph("Пользователь не выбрал ни одной фотографии погрузки.")

        p = doc.add_paragraph()
        run = p.add_run("ЗАМЕЧАНИЯ:")
        run.bold = True
        doc.add_paragraph("    1. Данный отчет касается указанного места и времени проведения инспекции.\n"
                          "    2. ИП «СЖС Минск» ООО не подтверждает точность, аккуратность и достоверность документов, предоставленных третьими сторонами.\n"
                          "    3. Данный отчет выпущен в соответствии с Общими условиями по оказанию инспекционных услуг.\n")


        from docx.enum.table import WD_TABLE_ALIGNMENT

        # Добавляем таблицу 1x2 без границ для подписей
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Убираем границы у таблицы
        tbl = table._tbl
        for cell in table._cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

        # Левый столбец (слева выравнено)
        left_texts = [
            "ПОДПИСАН И ВЫПУЩЕН В МИНСКЕ",
            "03 МАРТА 2025 ГОДА"
        ]

        for row_idx, text in enumerate(left_texts):
            cell = table.cell(row_idx, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Убираем межстрочные отступы
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)

        # Правый столбец (справа выравнено)
        right_texts = [
            "ОТ ИМЕНИ И ПО ПОРУЧЕНИЮ",
            "ИП «СЖС МИНСК» ООО"
        ]

        for row_idx, text in enumerate(right_texts):
            cell = table.cell(row_idx, 1)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Убираем межстрочные отступы
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)

        return doc, insp.inspection_date

    except Exception as e:
        print(f"Error in generating report for inspection {inspection_id}: {e}")
        raise
